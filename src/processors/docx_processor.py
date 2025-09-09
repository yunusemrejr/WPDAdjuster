"""
DOCX document processor using python-docx
"""

import logging
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_SECTION
import tempfile
import shutil

class DocxProcessor:
    """Processor for DOCX documents"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """
        Process DOCX document and extract properties
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing document properties
        """
        try:
            doc = Document(file_path)
            
            # Extract basic metadata
            properties = doc.core_properties
            
            # Extract page settings
            sections = doc.sections
            first_section = sections[0] if sections else None
            
            page_width = first_section.page_width.inches if first_section else 8.5
            page_height = first_section.page_height.inches if first_section else 11.0
            
            # Extract margins
            margins = {
                'top': first_section.top_margin.inches if first_section else 1.0,
                'bottom': first_section.bottom_margin.inches if first_section else 1.0,
                'left': first_section.left_margin.inches if first_section else 1.0,
                'right': first_section.right_margin.inches if first_section else 1.0
            }
            
            # Extract font information
            fonts_used = self._extract_fonts(doc)
            
            # Extract heading counts
            heading_counts = self._count_headings(doc)
            
            # Extract table count
            table_count = len(doc.tables)
            
            # Count paragraphs and words
            paragraph_count = len(doc.paragraphs)
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            
            return {
                'file_path': str(file_path),
                'filename': file_path.name,
                'file_type': 'docx',
                'title': properties.title or 'Untitled',
                'author': properties.author or 'Unknown',
                'page_count': len(sections),
                'word_count': word_count,
                'paragraph_count': paragraph_count,
                'page_dimensions': {
                    'width': page_width,
                    'height': page_height
                },
                'margins': margins,
                'fonts_used': fonts_used,
                'heading_counts': heading_counts,
                'table_count': table_count,
                'document': doc  # Keep document object for modifications
            }
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX document: {e}")
            raise ValueError(f"Failed to process DOCX document: {e}")
            
    def apply_modifications(self, document_data: Dict[str, Any], settings: Dict[str, Any]) -> Path:
        """
        Apply modification settings to DOCX document
        
        Args:
            document_data: Original document data
            settings: Modification settings
            
        Returns:
            Path to modified document
        """
        try:
            # Get original document
            doc = document_data['document']
            
            # Apply page size changes
            if 'page_size' in settings:
                self._apply_page_size(doc, settings['page_size'])
                
            # Apply margin changes
            if 'margins' in settings:
                self._apply_margins(doc, settings['margins'])
                
            # Apply font changes
            if 'font_settings' in settings:
                self._apply_font_settings(doc, settings['font_settings'])
                
            # Apply line spacing changes
            if 'line_spacing' in settings:
                self._apply_line_spacing(doc, settings['line_spacing'])
            
            # Save modified document
            output_path = self._get_output_path(document_data['file_path'])
            doc.save(output_path)
            
            self.logger.info(f"Modified DOCX saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error applying modifications to DOCX: {e}")
            raise ValueError(f"Failed to apply modifications: {e}")
            
    def _extract_fonts(self, doc: Document) -> Dict[str, int]:
        """Extract font usage statistics"""
        fonts = {}
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    font_name = run.font.name
                    fonts[font_name] = fonts.get(font_name, 0) + 1
                    
        return fonts
        
    def _count_headings(self, doc: Document) -> Dict[str, int]:
        """Count heading levels"""
        heading_counts = {'Heading 1': 0, 'Heading 2': 0, 'Heading 3': 0, 
                         'Heading 4': 0, 'Heading 5': 0, 'Heading 6': 0}
        
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            if style_name in heading_counts:
                heading_counts[style_name] += 1
                
        return heading_counts
        
    def _apply_page_size(self, doc: Document, page_size: Dict[str, Any]):
        """Apply page size changes"""
        width = page_size.get('width', 8.5)
        height = page_size.get('height', 11.0)
        
        for section in doc.sections:
            section.page_width = Inches(width)
            section.page_height = Inches(height)
            
    def _apply_margins(self, doc: Document, margins: Dict[str, float]):
        """Apply margin changes"""
        for section in doc.sections:
            if 'top' in margins:
                section.top_margin = Inches(margins['top'])
            if 'bottom' in margins:
                section.bottom_margin = Inches(margins['bottom'])
            if 'left' in margins:
                section.left_margin = Inches(margins['left'])
            if 'right' in margins:
                section.right_margin = Inches(margins['right'])
                
    def _apply_font_settings(self, doc: Document, font_settings: Dict[str, Any]):
        """Apply font changes"""
        font_family = font_settings.get('family')
        font_size = font_settings.get('size')
        
        if font_family or font_size:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if font_family:
                        run.font.name = font_family
                    if font_size:
                        run.font.size = Pt(font_size)
                        
    def _apply_line_spacing(self, doc: Document, line_spacing: float):
        """Apply line spacing changes"""
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = line_spacing
            
    def _get_output_path(self, original_path: str) -> Path:
        """Generate output path for modified document"""
        original = Path(original_path)
        output_name = f"{original.stem}_modified{original.suffix}"
        return original.parent / output_name
